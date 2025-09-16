def create_visualizations(df, numerical_cols, categorical_cols):
    """Create smart, business-focused visualizations based on data"""
    fig1, fig2 = None, None
    
    # Create visualization for the most business-relevant numerical column
    if numerical_cols:
        # Prioritize business metrics
        priority_cols = []
        for col in numerical_cols:
            col_lower = col.lower()
            if any(keyword in col_lower for keyword in ['satisfaction', 'rating', 'score', 'quality', 'performance']):
                priority_cols.append((col, 3))  # High priority
            elif any(keyword in col_lower for keyword in ['time', 'duration', 'response', 'wait']):
                priority_cols.append((col, 2))  # Medium priority
            else:
                priority_cols.append((col, 1))  # Low priority
        
        # Sort by priority and select the best column
        priority_cols.sort(key=lambda x: x[1], reverse=True)
        primary_num_col = priority_cols[0][0]
        
        fig1, ax1 = plt.subplots(figsize=(12, 7))
        
        # Remove any infinite or extremely large values
        clean_data = df[primary_num_col].replace([np.inf, -np.inf], np.nan).dropna()
        
        if len(clean_data) > 0:
            # Create histogram with better styling
            n, bins, patches = ax1.hist(clean_data, bins=min(30, int(np.sqrt(len(clean_data)))), 
                                      alpha=0.7, color='steelblue', edgecolor='darkblue', linewidth=0.8)
            
            # Color bars based on performance (if it's a rating/satisfaction metric)
            if any(keyword in primary_num_col.lower() for keyword in ['satisfaction', 'rating', 'score', 'quality']):
                max_val = clean_data.max()
                if max_val <= 5:  # 1-5 scale
                    for i, (patch, bin_val) in enumerate(zip(patches, bins[:-1])):
                        if bin_val < 2.5:
                            patch.set_facecolor('lightcoral')  # Poor performance
                        elif bin_val < 3.5:
                            patch.set_facecolor('gold')  # Average performance
                        else:
                            patch.set_facecolor('lightgreen')  # Good performance
            
            ax1.set_title(f'Distribution of {primary_num_col.replace("_", " ").title()}', 
                         fontsize=16, fontweight='bold', pad=20)
            ax1.set_xlabel(primary_num_col.replace("_", " ").title(), fontsize=12)
            ax1.set_ylabel('Frequency', fontsize=12)
            ax1.grid(axis='y', alpha=0.3, linestyle='--')
            
            # Add statistical annotations with business context
            mean_val = clean_data.mean()
            median_val = clean_data.median()
            std_val = clean_data.std()
            
            ax1.axvline(mean_val, color='red', linestyle='-', linewidth=2, label=f'Mean: {mean_val:.2f}')
            ax1.axvline(median_val, color='green', linestyle='-', linewidth=2, label=f'Median: {median_val:.2f}')
            
            # Add performance zones for rating metrics
            if any(keyword in primary_num_col.lower() for keyword in ['satisfaction', 'rating', 'score']):
                max_val = clean_data.max()
                if max_val <= 5:
                    ax1.axvspan(0, 2.5, alpha=0.1, color='red', label='Needs Improvement')
                    ax1.axvspan(2.5, 3.5, alpha=0.1, color='orange', label='Average')
                    ax1.axvspan(3.5, 5, alpha=0.1, color='green', label='Good')
            
            ax1.legend(loc='upper right')
            
            # Add summary text box
            textstr = f'n = {len(clean_data):,}\nStd Dev = {std_val:.2f}'
            props = dict(boxstyle='round', facecolor='wheat', alpha=0.5)
            ax1.text(0.02, 0.98, textstr, transform=ax1.transAxes, fontsize=10,
                    verticalalignment='top', bbox=props)
    
    # Create visualization for most business-relevant categorical column
    if categorical_cols:
        # Prioritize business-relevant categories
        priority_cats = []
        for col in categorical_cols:
            col_lower = col.lower()
            if any(keyword in col_lower for keyword in ['intent', 'category', 'type', 'status', 'priority']):
                priority_cats.append((col, 3))  # High priority
            elif any(keyword in col_lower for keyword in ['department', 'channel', 'source', 'method']):
                priority_cats.append((col, 2))  # Medium priority
            else:
                priority_cats.append((col, 1))  # Low priority
        
        if priority_cats:
            priority_cats.sort(key=lambda x: x[1], reverse=True)
            primary_cat_col = priority_cats[0][0]
            
            fig2, ax2 = plt.subplots(figsize=(12, 8))
            
            value_counts = df[primary_cat_col].value_counts().head(10)  # Top 10 categories
            
            if len(value_counts) > 0:
                # Create horizontal bar chart for better readability
                bars = ax2.barh(range(len(value_counts)), value_counts.values, 
                              color='lightcoral', alpha=0.8, edgecolor='darkred', linewidth=0.8)
                
                ax2.set_title(f'Top Categories: {primary_cat_col.replace("_", " ").title()}', 
                             fontsize=16, fontweight='bold', pad=20)
                ax2.set_xlabel('Frequency', fontsize=12)
                ax2.set_ylabel(primary_cat_col.replace("_", " ").title(), fontsize=12)
                ax2.set_yticks(range(len(value_counts)))
                ax2.set_yticklabels(value_counts.index)
                ax2.grid(axis='x', alpha=0.3, linestyle='--')
                
                # Add percentage labels on bars
                total_count = len(df)
                for i, (bar, value) in enumerate(zip(bars, value_counts.values)):
                    percentage = (value / total_count) * 100
                    ax2.text(bar.get_width() + total_count * 0.01, bar.get_y() + bar.get_height()/2,
                            f'{value:,} ({percentage:.1f}%)', 
                            ha='left', va='center', fontweight='bold')
                
                # Invert y-axis to show highest values at top
                ax2.invert_yaxis()
                
                # Add summary information
                unique_count = df[primary_cat_col].nunique()
                coverage = (value_counts.sum() / len(df)) * 100
                textstr = f'Total Categories: {unique_count}\nTop 10 Coverage: {coverage:.1f}%'
                props = dict(boxstyle='round', facecolor='lightblue', alpha=0.5)
                ax2.text(0.98, 0.02, textstr, transform=ax2.transAxes, fontsize=10,
                        verticalalignment='bottom', horizontalalignment='right', bbox=props)
            
            plt.tight_layout()
    
    return fig1, fig2
    
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
from docx import Document
from docx.shared import Inches
import io
import base64
from scipy.stats import pearsonr
import warnings
warnings.filterwarnings('ignore')

# Set page configuration
st.set_page_config(
    page_title="Advanced Data Analysis Tool",
    page_icon="📊",
    layout="wide"
)

def perform_analysis_and_generate_report(df):
    """
    Perform comprehensive data analysis and generate a Word report
    """
    # Initialize report buffer
    buffer = io.BytesIO()
    
    # Detect column types and filter out ID columns and timestamps
    numerical_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    categorical_cols = df.select_dtypes(include=['object', 'category']).columns.tolist()
    
    # Smart filtering of columns for analysis
    id_keywords = ['id', 'index', 'key', 'number', 'count', '_id']
    time_keywords = ['timestamp', 'date', 'time', 'created', 'updated']
    
    # Filter out ID and timestamp columns from numerical analysis
    meaningful_numerical = []
    for col in numerical_cols:
        col_lower = col.lower()
        is_id = any(keyword in col_lower for keyword in id_keywords)
        is_sequential = len(df[col].unique()) == len(df)  # Check if it's just sequential numbers
        if not is_id and not is_sequential:
            meaningful_numerical.append(col)
    
    # Filter categorical columns to focus on business-relevant ones
    meaningful_categorical = []
    for col in categorical_cols:
        col_lower = col.lower()
        is_timestamp = any(keyword in col_lower for keyword in time_keywords)
        uniqueness_ratio = df[col].nunique() / len(df)
        if not is_timestamp and uniqueness_ratio < 0.8:  # Not too granular
            meaningful_categorical.append(col)
    
    # Statistical Analysis
    stats_summary = {}
    correlation_insights = []
    
    # Analyze meaningful numerical columns
    numerical_stats = {}
    for col in meaningful_numerical:
        if df[col].notna().sum() > 0:  # Only analyze columns with data
            stats = {
                'mean': df[col].mean(),
                'median': df[col].median(),
                'std': df[col].std(),
                'min': df[col].min(),
                'max': df[col].max(),
                'top_5_values': df[col].nlargest(5).tolist()
            }
            numerical_stats[col] = stats
    
    # Analyze meaningful categorical columns
    categorical_stats = {}
    for col in meaningful_categorical:
        if df[col].notna().sum() > 0:  # Only analyze columns with data
            value_counts = df[col].value_counts().head(10)  # Get top 10 for better analysis
            categorical_stats[col] = {
                'top_categories': value_counts.to_dict(),
                'unique_count': df[col].nunique(),
                'most_frequent': value_counts.index[0] if len(value_counts) > 0 else None,
                'most_frequent_pct': (value_counts.iloc[0] / len(df)) * 100 if len(value_counts) > 0 else 0
            }
    
    # Correlation Analysis (only for meaningful numerical columns)
    strong_correlations = []
    if len(meaningful_numerical) >= 2:
        corr_matrix = df[meaningful_numerical].corr()
        for i in range(len(meaningful_numerical)):
            for j in range(i+1, len(meaningful_numerical)):
                correlation = corr_matrix.iloc[i, j]
                if abs(correlation) > 0.5 and not np.isnan(correlation):  # Lower threshold for more insights
                    strong_correlations.append({
                        'col1': meaningful_numerical[i],
                        'col2': meaningful_numerical[j],
                        'correlation': correlation
                    })
    
    # Generate Dynamic Narrative
    executive_summary = generate_executive_summary(df, numerical_stats, categorical_stats, strong_correlations)
    
    # Generate Recommendations
    recommendations = generate_recommendations(df, numerical_stats, categorical_stats, strong_correlations)
    
    # Generate Visualizations
    fig1, fig2 = create_visualizations(df, meaningful_numerical, meaningful_categorical)
    
    # Generate Word Document
    doc = Document()
    
    # Add Title
    title = doc.add_heading('Automated Data Analysis Report', 0)
    title.alignment = 1  # Center alignment
    
    # Add Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(executive_summary)
    
    # Add Key Findings
    doc.add_heading('Key Findings', level=1)
    
    # Numerical Statistics Table
    if numerical_stats:
        doc.add_heading('Numerical Analysis', level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Column'
        hdr_cells[1].text = 'Mean'
        hdr_cells[2].text = 'Median'
        hdr_cells[3].text = 'Std Dev'
        hdr_cells[4].text = 'Min'
        hdr_cells[5].text = 'Max'
        
        for col, stats in numerical_stats.items():
            row_cells = table.add_row().cells
            row_cells[0].text = col
            row_cells[1].text = f"{stats['mean']:.2f}"
            row_cells[2].text = f"{stats['median']:.2f}"
            row_cells[3].text = f"{stats['std']:.2f}"
            row_cells[4].text = f"{stats['min']:.2f}"
            row_cells[5].text = f"{stats['max']:.2f}"
    
    # Categorical Statistics
    if categorical_stats:
        doc.add_heading('Categorical Analysis', level=2)
        for col, stats in categorical_stats.items():
            doc.add_paragraph(f"Column '{col}': {stats['unique_count']} unique values")
            doc.add_paragraph(f"Most frequent value: {stats['most_frequent']}")
    
    # Add Recommendations
    doc.add_heading('Recommendations', level=1)
    for rec in recommendations:
        doc.add_paragraph(rec, style='List Bullet')
    
    # Add Visualizations
    doc.add_heading('Data Visualizations', level=1)
    
    # Save plots and add to document
    if fig1 is not None:
        img_buffer1 = io.BytesIO()
        fig1.savefig(img_buffer1, format='png', dpi=300, bbox_inches='tight')
        img_buffer1.seek(0)
        doc.add_paragraph('Distribution Analysis:')
        doc.add_picture(img_buffer1, width=Inches(6))
        plt.close(fig1)
    
    if fig2 is not None:
        img_buffer2 = io.BytesIO()
        fig2.savefig(img_buffer2, format='png', dpi=300, bbox_inches='tight')
        img_buffer2.seek(0)
        doc.add_paragraph('Category Frequency Analysis:')
        doc.add_picture(img_buffer2, width=Inches(6))
        plt.close(fig2)
    
    # Save document to buffer
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def generate_executive_summary(df, numerical_stats, categorical_stats, strong_correlations):
    """Generate intelligent business-focused narrative based on analysis results"""
    summary = f"This report analyzes a dataset containing {len(df):,} records across {len(df.columns)} dimensions. "
    
    # Identify key business metrics
    business_metrics = {}
    for col, stats in numerical_stats.items():
        col_lower = col.lower()
        if any(word in col_lower for word in ['satisfaction', 'rating', 'score', 'quality', 'performance']):
            business_metrics[col] = stats
    
    # Focus on business-critical insights
    if business_metrics:
        for col, stats in business_metrics.items():
            mean_val = stats['mean']
            scale_max = stats['max']
            
            # Determine likely scale and performance level
            if scale_max <= 5:  # 1-5 scale
                if mean_val < 2.5:
                    performance = "critically low"
                elif mean_val < 3.5:
                    performance = "below average"
                elif mean_val < 4.0:
                    performance = "moderate"
                else:
                    performance = "strong"
                summary += f"The {col.replace('_', ' ')} averages {mean_val:.1f} out of {scale_max}, indicating {performance} performance. "
            
            elif scale_max <= 10:  # 1-10 scale
                if mean_val < 5:
                    performance = "below expectations"
                elif mean_val < 7:
                    performance = "moderate"
                else:
                    performance = "above average"
                summary += f"The {col.replace('_', ' ')} averages {mean_val:.1f} out of {scale_max:.0f}, showing {performance} results. "
    
    # Add operational insights
    time_based_cols = [col for col in df.columns if any(word in col.lower() for word in ['time', 'duration', 'response', 'wait'])]
    if time_based_cols and time_based_cols[0] in numerical_stats:
        time_col = time_based_cols[0]
        time_stats = numerical_stats[time_col]
        avg_time = time_stats['mean']
        
        if 'response' in time_col.lower() or 'wait' in time_col.lower():
            if avg_time > 60:  # Assuming seconds
                summary += f"Average {time_col.replace('_', ' ')} is {avg_time:.0f} seconds, which may impact user experience. "
            else:
                summary += f"Response times average {avg_time:.1f} seconds, indicating efficient service delivery. "
    
    # Category insights with business context
    priority_categories = ['intent', 'category', 'type', 'status', 'priority', 'department', 'channel']
    for col, stats in categorical_stats.items():
        col_lower = col.lower()
        if any(priority in col_lower for priority in priority_categories):
            most_frequent = stats['most_frequent']
            percentage = stats['most_frequent_pct']
            summary += f"The most common {col.replace('_', ' ')} is '{most_frequent}' ({percentage:.1f}% of cases), "
            
            # Add business interpretation
            if 'support' in most_frequent.lower() or 'technical' in most_frequent.lower():
                summary += "suggesting a focus on technical assistance needs. "
            elif 'sales' in most_frequent.lower() or 'billing' in most_frequent.lower():
                summary += "indicating commercial inquiry patterns. "
            else:
                summary += f"representing the primary use case pattern. "
            break
    
    # Correlation insights with business meaning
    if strong_correlations:
        strongest_corr = max(strong_correlations, key=lambda x: abs(x['correlation']))
        correlation_strength = "strong" if abs(strongest_corr['correlation']) > 0.7 else "moderate"
        correlation_direction = "positive" if strongest_corr['correlation'] > 0 else "negative"
        
        summary += f"A {correlation_strength} {correlation_direction} relationship exists between {strongest_corr['col1'].replace('_', ' ')} and {strongest_corr['col2'].replace('_', ' ')} (r={strongest_corr['correlation']:.2f}), "
        
        # Business interpretation of correlation
        if 'satisfaction' in strongest_corr['col1'].lower() or 'satisfaction' in strongest_corr['col2'].lower():
            summary += "highlighting factors that directly impact customer satisfaction. "
        elif 'time' in strongest_corr['col1'].lower() or 'time' in strongest_corr['col2'].lower():
            summary += "indicating operational efficiency relationships. "
        else:
            summary += "revealing important operational dependencies. "
    
    return summary

def generate_recommendations(df, numerical_stats, categorical_stats, strong_correlations):
    """Generate intelligent, actionable business recommendations based on analysis"""
    recommendations = []
    
    # Priority 1: Address critical performance issues
    critical_issues = []
    improvement_opportunities = []
    operational_insights = []
    
    # Analyze business metrics for critical issues
    for col, stats in numerical_stats.items():
        col_lower = col.lower()
        mean_val = stats['mean']
        max_val = stats['max']
        
        # Customer satisfaction and quality metrics
        if any(keyword in col_lower for keyword in ['satisfaction', 'rating', 'quality', 'score']):
            if max_val <= 5:  # 1-5 scale
                if mean_val < 2.5:
                    critical_issues.append(f"🚨 CRITICAL: {col.replace('_', ' ').title()} is critically low at {mean_val:.1f}/5. Immediate intervention required to prevent customer churn.")
                elif mean_val < 3.5:
                    improvement_opportunities.append(f"📈 IMPROVE: {col.replace('_', ' ').title()} at {mean_val:.1f}/5 has significant room for improvement. Target: 4.0+ through focused quality initiatives.")
                elif mean_val >= 4.0:
                    operational_insights.append(f"✅ MAINTAIN: {col.replace('_', ' ').title()} is performing well at {mean_val:.1f}/5. Identify and replicate success factors.")
            
            elif max_val <= 10:  # 1-10 scale
                if mean_val < 5:
                    critical_issues.append(f"🚨 CRITICAL: {col.replace('_', ' ').title()} at {mean_val:.1f}/10 requires immediate attention.")
                elif mean_val < 7:
                    improvement_opportunities.append(f"📈 IMPROVE: {col.replace('_', ' ').title()} at {mean_val:.1f}/10 - target 7.5+ for competitive performance.")
        
        # Response time and operational metrics
        elif any(keyword in col_lower for keyword in ['response', 'wait', 'duration', 'time']):
            if mean_val > 60:  # Assuming seconds - slow response
                critical_issues.append(f"⏱️ OPTIMIZE: Average {col.replace('_', ' ')} of {mean_val:.0f} seconds is too high. Target: <30 seconds for improved user experience.")
            elif mean_val > 30:
                improvement_opportunities.append(f"🎯 ENHANCE: {col.replace('_', ' ')} averaging {mean_val:.0f} seconds could be optimized for better efficiency.")
        
        # Variability analysis - only for business-relevant metrics
        cv = stats['std'] / stats['mean'] if stats['mean'] != 0 else 0
        if cv > 0.4 and any(keyword in col_lower for keyword in ['satisfaction', 'quality', 'performance', 'rating']):
            operational_insights.append(f"📊 STANDARDIZE: High variability in {col.replace('_', ' ')} (CV: {cv:.1f}) suggests inconsistent delivery. Implement quality control measures.")
    
    # Category-based insights and recommendations
    for col, stats in categorical_stats.items():
        col_lower = col.lower()
        most_frequent = stats['most_frequent']
        frequency_pct = stats['most_frequent_pct']
        
        # Intent/Category analysis
        if 'intent' in col_lower or 'category' in col_lower:
            if 'support' in most_frequent.lower() and frequency_pct > 40:
                improvement_opportunities.append(f"🛠️ PROACTIVE: {frequency_pct:.1f}% of cases are '{most_frequent}' - consider proactive measures to reduce support volume through better documentation or product improvements.")
            elif 'complaint' in most_frequent.lower():
                critical_issues.append(f"🚨 ATTENTION: High volume of '{most_frequent}' cases suggests systematic issues requiring investigation.")
        
        # Channel/Department analysis
        elif any(keyword in col_lower for keyword in ['channel', 'department', 'source']):
            if frequency_pct > 60:
                operational_insights.append(f"📍 RESOURCE: {frequency_pct:.1f}% of activity comes through '{most_frequent}' - ensure this channel has adequate resources and optimization.")
        
        # Status/Performance analysis
        elif any(keyword in col_lower for keyword in ['status', 'performance', 'outcome']):
            if 'poor' in most_frequent.lower() or 'failed' in most_frequent.lower() or 'needs improvement' in most_frequent.lower():
                critical_issues.append(f"❌ RESOLVE: High frequency of '{most_frequent}' status indicates process failures requiring immediate attention.")
    
    # Correlation-based recommendations
    for corr in strong_correlations:
        col1, col2, corr_val = corr['col1'], corr['col2'], corr['correlation']
        
        if abs(corr_val) > 0.7:  # Strong correlation
            if 'satisfaction' in col1.lower() or 'satisfaction' in col2.lower():
                if corr_val > 0:
                    improvement_opportunities.append(f"💡 LEVERAGE: Strong positive correlation between {col1.replace('_', ' ')} and {col2.replace('_', ' ')} (r={corr_val:.2f}). Improving one will boost the other.")
                else:
                    critical_issues.append(f"⚠️ MONITOR: Negative correlation between {col1.replace('_', ' ')} and {col2.replace('_', ' ')} (r={corr_val:.2f}). Balance these competing factors carefully.")
            
            elif 'time' in col1.lower() or 'time' in col2.lower():
                if corr_val < 0:
                    operational_insights.append(f"⚡ EFFICIENCY: Reducing {col1.replace('_', ' ')} correlates with improved {col2.replace('_', ' ')} - focus on speed optimization.")
    
    # Compile recommendations in priority order
    final_recommendations = []
    
    # Add critical issues first
    final_recommendations.extend(critical_issues)
    
    # Add improvement opportunities
    final_recommendations.extend(improvement_opportunities)
    
    # Add operational insights
    final_recommendations.extend(operational_insights)
    
    # Add data quality recommendations if needed
    missing_data = df.isnull().sum().sum()
    if missing_data > len(df) * 0.05:  # More than 5% missing data
        final_recommendations.append(f"🔍 DATA QUALITY: {missing_data:,} missing values detected. Implement data collection improvements to enhance analysis reliability.")
    
    # Add sample size recommendations
    if len(df) < 100:
        final_recommendations.append("📊 SAMPLE SIZE: Consider collecting more data points for more robust statistical analysis and reliable insights.")
    
    # Default recommendation if nothing specific found
    if not final_recommendations:
        final_recommendations.append("✅ MONITORING: Current performance appears stable. Continue regular monitoring and consider establishing KPI benchmarks for ongoing measurement.")
    
    return final_recommendations

def create_visualizations(df, numerical_cols, categorical_cols):
    """Create dynamic visualizations based on data"""
    fig1, fig2 = None, None
    
    # Create histogram for primary numerical column
    if numerical_cols:
        primary_num_col = numerical_cols[0]  # Use first numerical column
        fig1, ax1 = plt.subplots(figsize=(10, 6))
        
        # Remove any infinite or extremely large values
        clean_data = df[primary_num_col].replace([np.inf, -np.inf], np.nan).dropna()
        
        if len(clean_data) > 0:
            ax1.hist(clean_data, bins=30, alpha=0.7, color='skyblue', edgecolor='black')
            ax1.set_title(f'Distribution of {primary_num_col}', fontsize=14, fontweight='bold')
            ax1.set_xlabel(primary_num_col, fontsize=12)
            ax1.set_ylabel('Frequency', fontsize=12)
            ax1.grid(axis='y', alpha=0.3)
            
            # Add statistical annotations
            mean_val = clean_data.mean()
            median_val = clean_data.median()
            ax1.axvline(mean_val, color='red', linestyle='--', label=f'Mean: {mean_val:.2f}')
            ax1.axvline(median_val, color='green', linestyle='--', label=f'Median: {median_val:.2f}')
            ax1.legend()
    
    # Create bar chart for primary categorical column
    if categorical_cols:
        primary_cat_col = categorical_cols[0]  # Use first categorical column
        fig2, ax2 = plt.subplots(figsize=(10, 6))
        
        value_counts = df[primary_cat_col].value_counts().head(10)  # Top 10 categories
        
        if len(value_counts) > 0:
            bars = ax2.bar(range(len(value_counts)), value_counts.values, 
                          color='lightcoral', alpha=0.8, edgecolor='black')
            ax2.set_title(f'Top Categories in {primary_cat_col}', fontsize=14, fontweight='bold')
            ax2.set_xlabel(primary_cat_col, fontsize=12)
            ax2.set_ylabel('Frequency', fontsize=12)
            ax2.set_xticks(range(len(value_counts)))
            ax2.set_xticklabels(value_counts.index, rotation=45, ha='right')
            ax2.grid(axis='y', alpha=0.3)
            
            # Add value labels on bars
            for bar, value in zip(bars, value_counts.values):
                ax2.text(bar.get_x() + bar.get_width()/2, bar.get_height() + value*0.01,
                        str(value), ha='center', va='bottom')
        
        plt.tight_layout()
    
    return fig1, fig2

def main():
    """Main Streamlit application"""
    # Title and description
    st.title("📊 Advanced Data Analysis Tool")
    st.markdown("Upload a CSV file to generate comprehensive analysis and downloadable report")
    
    # Sidebar with instructions
    with st.sidebar:
        st.header("Instructions")
        st.markdown("""
        1. Upload your CSV file
        2. Click 'Generate Report'
        3. View the analysis results
        4. Download the Word document report
        
        **Features:**
        - Automatic statistical analysis
        - Dynamic narrative generation
        - Actionable recommendations
        - Professional visualizations
        - Downloadable Word report
        """)
    
    # File uploader
    uploaded_file = st.file_uploader("Choose a CSV file", type="csv")
    
    if uploaded_file is not None:
        try:
            # Read the CSV file
            df = pd.read_csv(uploaded_file)
            
            # Display basic info about the dataset
            st.subheader("Dataset Overview")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Rows", len(df))
            with col2:
                st.metric("Columns", len(df.columns))
            with col3:
                st.metric("Missing Values", df.isnull().sum().sum())
            
            # Show first few rows
            st.subheader("Data Preview")
            st.dataframe(df.head(10))
            
            # Generate Report Button
            if st.button("Generate Report", type="primary"):
                with st.spinner("Analyzing data and generating report..."):
                    try:
                        # Perform analysis and generate report
                        report_buffer = perform_analysis_and_generate_report(df)
                        
                        st.success("Report generated successfully!")
                        
                        # Provide download button
                        st.download_button(
                            label="📄 Download Analysis Report (Word Document)",
                            data=report_buffer.getvalue(),
                            file_name="data_analysis_report.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        # Display some key insights in the app
                        st.subheader("Key Insights Preview")
                        
                        # Smart stats display
                        meaningful_numerical = [col for col in numerical_cols 
                                              if not any(keyword in col.lower() for keyword in ['id', 'index', 'number', 'count'])
                                              and len(df[col].unique()) < len(df)]  # Not sequential IDs
                        meaningful_categorical = [col for col in categorical_cols 
                                                if not any(keyword in col.lower() for keyword in ['timestamp', 'date', 'time', 'created'])
                                                and df[col].nunique() / len(df) < 0.8]  # Not too granular
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            st.info(f"**Key Numerical Metrics:** {len(meaningful_numerical)}")
                            if meaningful_numerical:
                                # Show business metrics first
                                business_metrics = [col for col in meaningful_numerical 
                                                  if any(word in col.lower() for word in ['satisfaction', 'rating', 'score', 'quality', 'performance'])]
                                if business_metrics:
                                    for col in business_metrics[:3]:
                                        avg_val = df[col].mean()
                                        max_val = df[col].max()
                                        if max_val <= 5:
                                            status = "🔴" if avg_val < 2.5 else "🟡" if avg_val < 3.5 else "🟢"
                                        else:
                                            status = "🟢"
                                        st.write(f"{status} **{col.replace('_', ' ').title()}:** {avg_val:.1f}/{max_val}")
                                else:
                                    st.write("• " + "\n• ".join(meaningful_numerical[:5]))
                        
                        with col2:
                            st.info(f"**Key Categories:** {len(meaningful_categorical)}")
                            if meaningful_categorical:
                                # Show most important categorical insights
                                for col in meaningful_categorical[:3]:
                                    most_frequent = df[col].value_counts().index[0]
                                    frequency = df[col].value_counts().iloc[0]
                                    percentage = (frequency / len(df)) * 100
                                    st.write(f"📊 **{col.replace('_', ' ').title()}:** {most_frequent} ({percentage:.1f}%)")
                            else:
                                st.write("• " + "\n• ".join(meaningful_categorical[:5]))
                        
                        # Show critical alerts if any
                        critical_metrics = []
                        for col in meaningful_numerical:
                            if any(word in col.lower() for word in ['satisfaction', 'rating', 'quality']):
                                avg_val = df[col].mean()
                                max_val = df[col].max()
                                if max_val <= 5 and avg_val < 2.5:
                                    critical_metrics.append(f"🚨 **{col.replace('_', ' ').title()}** is critically low at {avg_val:.1f}/5")
                                elif max_val <= 5 and avg_val < 3.5:
                                    critical_metrics.append(f"⚠️ **{col.replace('_', ' ').title()}** needs improvement at {avg_val:.1f}/5")
                        
                        if critical_metrics:
                            st.warning("**Immediate Attention Required:**")
                            for metric in critical_metrics:
                                st.write(metric)
                        
                    except Exception as e:
                        st.error(f"Error generating report: {str(e)}")
                        st.info("Please ensure your CSV file is properly formatted and contains analyzable data.")
        
        except Exception as e:
            st.error(f"Error reading CSV file: {str(e)}")
            st.info("Please check that your file is a valid CSV format.")
    
    else:
        st.info("Please upload a CSV file to begin analysis.")
        
        # Show example of what the tool can do
        with st.expander("See example analysis"):
            st.markdown("""
            **This tool will automatically:**
            
            🔍 **Analyze your data:**
            - Detect numerical and categorical columns
            - Calculate statistical measures (mean, median, std dev)
            - Identify correlations and outliers
            - Generate frequency distributions
            
            📝 **Create narratives:**
            - Dynamic executive summary
            - Data-driven insights
            - Statistical interpretations
            
            💡 **Provide recommendations:**
            - Actionable business insights
            - Process improvements
            - Data quality suggestions
            
            📊 **Generate visualizations:**
            - Distribution histograms
            - Category frequency charts
            - Professional styling
            
            📄 **Produce Word report:**
            - Executive summary
            - Statistical tables
            - Visualizations
            - Recommendations
            """)

if __name__ == "__main__":
    main()
